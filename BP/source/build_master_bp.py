from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "育见商业计划书.md"
OUT_PATH = ROOT / "育见商业计划书.docx"

PAGE_W_DXA = 12240
PAGE_H_DXA = 15840
MARGIN_DXA = 1440
CONTENT_W_DXA = 9360

INK = "202633"
GREEN = "5F8F48"
GREEN_LIGHT = "EEF7E8"
GREEN_PALE = "F6FAF2"
MUTED = "666C76"
LINE = "D9E2D4"
WHITE = "FFFFFF"
CAUTION = "8A6414"
RISK = "9B1C1C"

BODY_FONT = "PingFang SC"
LATIN_FONT = "Arial"


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name=BODY_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def configure_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Mark the first row so screen readers and Word can identify/repeat it.
    first_row_pr = table.rows[0]._tr.get_or_add_trPr()
    header = first_row_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        first_row_pr.append(header)
    header.set(qn("w:val"), "true")


def set_table_borders(table, color=LINE, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, 9, MUTED)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, 10, MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id, num_fmt, text, left, hanging):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num = OxmlElement("w:abstractNumId")
        abstract_num.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num)
        numbering.append(num)

    add_abstract(20, "bullet", "•", 540, 280)
    add_num(20, 20)
    add_abstract(21, "decimal", "%1.", 540, 280)
    add_num(21, 21)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def parse_inline(paragraph, text, base_size=11, color=INK):
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`|https?://\S+)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, base_size, color, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, base_size - 0.5, GREEN, name="Menlo")
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), GREEN_PALE)
            run._r.get_or_add_rPr().append(shading)
        elif token.startswith("http://") or token.startswith("https://"):
            add_hyperlink(paragraph, token, token.rstrip("。"))
        else:
            run = paragraph.add_run(token)
            set_run_font(run, base_size, color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in (
        ("Heading 1", 18, 18, 10, GREEN),
        ("Heading 2", 14, 12, 6, GREEN),
        ("Heading 3", 12, 8, 4, INK),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def set_page_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)


def set_running_furniture(section):
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("育见 AI 家庭教育理解系统  |  Master BP Document")
    set_run_font(run, 8.5, MUTED, bold=False)

    footer = section.footer
    add_page_field(footer.paragraphs[0])

    # The cover is intentionally free of running furniture.
    section.first_page_header.paragraphs[0].clear()
    section.first_page_footer.paragraphs[0].clear()


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MASTER BP DOCUMENT")
    set_run_font(r, 11, GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("育见")
    set_run_font(r, 36, GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("AI 家庭教育理解系统")
    set_run_font(r, 22, INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("通过 AI 长期理解家庭与孩子，辅助家庭成长决策")
    set_run_font(r, 12.5, MUTED)

    table = doc.add_table(rows=3, cols=2)
    configure_table_geometry(table, [2800, 6560])
    set_table_borders(table, color=WHITE, size="0")
    data = (
        ("项目阶段", "产品验证与公益试点期"),
        ("文档版本", "V1.0 · 2026 年 7 月"),
        ("用途", "投资机构 / 高校创业比赛 / 产业合作 / 政府及基金会项目"),
    )
    for i, (label, value) in enumerate(data):
        set_cell_shading(table.cell(i, 0), GREEN_PALE)
        set_cell_shading(table.cell(i, 1), WHITE)
        p1 = table.cell(i, 0).paragraphs[0]
        p2 = table.cell(i, 1).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        parse_inline(p1, label, 9.5, GREEN)
        parse_inline(p2, value, 9.5, INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("可编辑商业计划书母稿")
    set_run_font(r, 10, MUTED)
    doc.add_page_break()


def add_document_note(doc):
    table = doc.add_table(rows=1, cols=1)
    configure_table_geometry(table, [9360])
    set_table_borders(table, color=LINE, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, GREEN_PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    parse_inline(
        p,
        "使用说明：本母稿用于拆解投资机构版 PPT、创业比赛申报书、产业合作介绍册及政府/基金会材料。事实、推断与规划已尽量分开；所有占位项在对外使用前必须补齐或删除。",
        10,
        INK,
    )


def add_static_contents(doc):
    entries = [
        ("执行摘要", "项目价值、阶段判断与核心结论"),
        ("一、项目概览", "定位、边界、目标用户与核心价值"),
        ("二、用户与问题", "家庭场景、痛点与价值假设"),
        ("三、行业背景与机会", "政策、用户规模与 AI 机会窗口"),
        ("四、产品体系", "用户旅程、四类入口与产品页面"),
        ("五、核心技术体系", "FamilyModel、事实网络、长期记忆与推理"),
        ("六、竞争与差异化", "通用 AI、传统咨询与育见对比"),
        ("七、产品进展与验证", "当前能力、运行状态与待验证事项"),
        ("八、用户验证与增长", "种子用户、学校社区与增长路径"),
        ("九、商业模式", "家庭订阅、机构服务及验证顺序"),
        ("十、技术与产品路线图", "短中长期建设重点"),
        ("十一、团队", "能力结构、缺失信息与补强计划"),
        ("十二、数据安全与合规", "隐私、未成年人保护与治理机制"),
        ("十三、融资规划", "资金用途框架与融资前置条件"),
        ("十四、风险与应对", "产品、技术、增长、商业和合规风险"),
        ("十五、未来 12 个月里程碑", "验证目标与阶段性判断标准"),
        ("十六、结论", "项目判断与下一阶段重点"),
        ("附录 A—C", "信息缺口、素材清单与资料来源"),
    ]
    row_count = (len(entries) + 1) // 2
    table = doc.add_table(rows=row_count, cols=2)
    configure_table_geometry(table, [4680, 4680])
    set_table_borders(table, color=LINE, size="4")
    for idx, (section, summary) in enumerate(entries):
        row_idx = idx % row_count
        col_idx = idx // row_count
        cell = table.cell(row_idx, col_idx)
        if row_idx % 2:
            set_cell_shading(cell, GREEN_PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(1)
        parse_inline(p, section, 8.8, GREEN)
        for run in p.runs:
            run.bold = True
        p.add_run().add_break()
        parse_inline(p, summary, 8.2, INK)


def add_image(doc, rel_path, caption):
    path = ROOT / rel_path
    if not path.exists():
        p = doc.add_paragraph()
        parse_inline(p, f"[图片缺失：{rel_path}]", 10, RISK)
        return
    is_mobile = path.suffix.lower() in {".jpg", ".jpeg"} and "产品截图" in str(path)
    width = Inches(2.65 if is_mobile else 6.15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=width)
    for doc_pr in run._r.xpath(".//wp:docPr"):
        doc_pr.set("descr", caption)
        doc_pr.set("title", caption)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    r = cap.add_run(f"图：{caption}")
    set_run_font(r, 9, MUTED, italic=True)


def add_markdown_table(doc, rows):
    col_count = max(len(row) for row in rows)
    if col_count == 2:
        widths = [2700, 6660]
    elif col_count == 3:
        widths = [1800, 3780, 3780]
    elif col_count == 4:
        widths = [1600, 2450, 2450, 2860]
    elif col_count == 5:
        widths = [1300, 1550, 2150, 2350, 2010]
    else:
        base = CONTENT_W_DXA // col_count
        widths = [base] * col_count
        widths[-1] += CONTENT_W_DXA - sum(widths)

    table = doc.add_table(rows=len(rows), cols=col_count)
    configure_table_geometry(table, widths)
    set_table_borders(table)

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if r_idx == 0:
                set_cell_shading(cell, GREEN)
                parse_inline(p, text, 9, WHITE)
                for run in p.runs:
                    run.bold = True
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, GREEN_PALE)
                parse_inline(p, text, 9, INK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [x.strip() for x in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part or "") for part in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_numbering_definitions(doc)
    for section in doc.sections:
        set_page_geometry(section)
        set_running_furniture(section)

    add_cover(doc)
    add_document_note(doc)

    toc_title = doc.add_paragraph(style="Heading 1")
    toc_title.add_run("目录")
    intro = doc.add_paragraph()
    parse_inline(intro, "以下结构可直接拆解为路演 PPT、申报书或合作材料；Word 的导航窗格可按标题快速定位。", 10, MUTED)
    add_static_contents(doc)
    doc.add_page_break()

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    # Skip the markdown cover block through the first separator.
    first_sep = next(i for i, line in enumerate(lines) if line.strip() == "---")
    i = first_sep + 1

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        if line == "---":
            doc.add_page_break()
            i += 1
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", line)
        if image_match:
            add_image(doc, image_match.group(2), image_match.group(1))
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                add_markdown_table(doc, rows)
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            level = min(level, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            parse_inline(p, text, {1: 18, 2: 14, 3: 12}[level], GREEN if level < 3 else INK)
            continue_idx = i + 1
            i = continue_idx
            continue

        if line.startswith(">"):
            table = doc.add_table(rows=1, cols=1)
            configure_table_geometry(table, [9360])
            set_table_borders(table, color=LINE, size="6")
            cell = table.cell(0, 0)
            set_cell_shading(cell, GREEN_PALE)
            p = cell.paragraphs[0]
            parse_inline(p, line[1:].strip(), 10.5, INK)
            i += 1
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            table = doc.add_table(rows=1, cols=1)
            configure_table_geometry(table, [9360])
            set_table_borders(table, color=LINE, size="4")
            cell = table.cell(0, 0)
            set_cell_shading(cell, "F7F8F6")
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run("\n".join(code_lines))
            set_run_font(r, 9.5, INK, name="Menlo")
            continue

        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph()
            apply_num(p, 21)
            parse_inline(p, text, 11, INK)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph()
            apply_num(p, 20)
            parse_inline(p, line[2:], 11, INK)
            i += 1
            continue

        if re.fullmatch(r"https?://\S+", line):
            p = doc.add_paragraph()
            add_hyperlink(p, "查看官方原文", line)
            i += 1
            continue

        p = doc.add_paragraph()
        parse_inline(p, line, 11, INK)
        i += 1

    doc.core_properties.title = "育见 AI 家庭教育理解系统 - 商业计划书母稿"
    doc.core_properties.subject = "Master BP Document"
    doc.core_properties.author = "育见项目团队"
    doc.core_properties.keywords = "育见, FamilyModel, SecondMe, 家庭教育, AI"

    # Ensure every section follows the same geometry/furniture after page breaks.
    for section in doc.sections:
        set_page_geometry(section)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
