#!/usr/bin/env python3
"""Generate 育见 Master BP Word document with embedded images."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).parent
MD = BASE / "育见商业计划书母稿.md"
OUT = BASE / "育见商业计划书母稿.docx"


def set_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.paragraph_format.line_spacing = 1.35
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "PingFang SC"
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_cover(doc: Document) -> None:
    cover_img = BASE / "images/产品截图/00-封面.png"
    if cover_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(cover_img), width=Cm(12))

    for _ in range(2):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("育见 AI 家庭教育理解系统")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("项目商业说明书（Master BP）")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x4A, 0x4A, 0x68)

    doc.add_paragraph()
    for line in [
        "文档版本：Master V1.0",
        "编制日期：2026 年 7 月",
        "产品体验：https://yujian.yihe.site",
        "清华大学 iCenter 支持 · AI 教育公益试点",
        "",
        "本文档为母稿，可拆解为投资 PPT、创业申报、产业合作等材料",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x80)

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_rich_paragraph(doc: Document, text: str, *, bullet: bool = False, number: bool = False) -> None:
    style = "List Bullet" if bullet else ("List Number" if number else None)
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            p.add_run(part)


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    img_path = (BASE / rel_path).resolve()
    if not img_path.exists():
        p = doc.add_paragraph(f"[图片缺失：{rel_path}]")
        p.runs[0].font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(13.5))
    cap = doc.add_paragraph(alt)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x80)
    doc.add_paragraph()


def parse_md_to_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    set_doc_styles(doc)
    add_cover(doc)

    lines = MD.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    table_buf: list[list[str]] = []
    in_table = False
    skip_until_body = True

    while i < len(lines):
        line = lines[i]
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())

        if skip_until_body:
            if line.startswith("# 执行摘要"):
                skip_until_body = False
                doc.add_heading("执行摘要", level=1)
            i += 1
            continue

        if img_match:
            if in_table and table_buf:
                add_table(doc, table_buf[0], table_buf[1:])
                table_buf, in_table = [], False
            add_image(doc, img_match.group(1), img_match.group(2))
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs:
                r.font.name = "Menlo"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            if not in_table:
                in_table = True
                table_buf = []
            if not re.match(r"^\|[-:\s|]+\|$", line):
                table_buf.append([c.strip() for c in line.strip("|").split("|")])
            i += 1
            continue
        elif in_table:
            if table_buf:
                add_table(doc, table_buf[0], table_buf[1:])
            table_buf, in_table = [], False

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            p = doc.add_paragraph()
            r = p.add_run(line[5:].strip())
            r.bold = True
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.paragraph_format.left_indent = Cm(0.4)
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x70)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip().startswith("- "):
            add_rich_paragraph(doc, line.strip()[2:], bullet=True)
        elif re.match(r"^\d+\.\s", line.strip()):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s", "", line.strip()), number=True)
        elif line.strip():
            add_rich_paragraph(doc, line)
        i += 1

    if in_table and table_buf:
        add_table(doc, table_buf[0], table_buf[1:])

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    parse_md_to_docx()
